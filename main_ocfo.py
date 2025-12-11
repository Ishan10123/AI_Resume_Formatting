# main_ocfo.py
import os
import re
import json
import shutil
import subprocess
import hashlib
from typing import Any, Dict, List, Optional, Tuple
from fastapi import FastAPI, File, UploadFile, Form, Query
from fastapi.responses import JSONResponse, HTMLResponse, FileResponse
from docx.table import Table
# ollama cloud client
from ollama import Client
from typing import Dict, Any
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table as _Table
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

import logging
import ast
import json as _json
import re as _re

try:
    import fitz  # PyMuPDF
except Exception as e:
    raise ImportError(
        "PyMuPDF is required (install with: python -m pip install pymupdf). "
        f"Import failed: {e}"
    )

logger = logging.getLogger(__name__)
# -------------------------
# Configuration
# -------------------------
os.environ["PATH"] += ":/usr/local/bin"

# Local sample screenshot path (use this path as the url when invoking tools)
SAMPLE_SCREENSHOT_PATH = "/mnt/data/Screenshot 2025-11-22 at 11.56.17 PM.png"

# Required env:
OLLAMA_API_KEY = os.environ.get("bce08f66335b43ed8feeca10d0b64876.g7V-_vkd4xJxkpbfWgG5YJKg")
if not OLLAMA_API_KEY:
    # we won't raise at import time to allow local linting; endpoints will error if missing.
    pass

# Cloud api host (Ollama cloud)
OLLAMA_API_HOST = os.environ.get("OLLAMA_CLOUD_URL", "https://ollama.com")
# chosen model
CLOUD_MODEL = os.environ.get("OLLAMA_CLOUD_MODEL", "gpt-oss:120b-cloud")

OUTPUT_DIR = "formatted_resumes"
TEMPLATES_DIR = "templates"
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)
OCFO_MODE = True

# create Ollama client once (will use Authorization header)
def make_ollama_client():
    headers = {"Authorization": "Bearer bce08f66335b43ed8feeca10d0b64876.g7V-_vkd4xJxkpbfWgG5YJKg"} if OLLAMA_API_KEY else {}
    return Client(host=OLLAMA_API_HOST, headers=headers)

# -------------------------
# FastAPI app
# -------------------------
app = FastAPI(
    title="Smart Resume Formatter (Ollama Cloud)",
    version="1.0",
    description="Use Ollama Cloud to extract resumes into a universal JSON schema and render DOCX templates."
)

# -------------------------
# Helpers: file convert/extract
# -------------------------
def convert_doc_to_docx(input_path: str) -> str:
    out_dir = os.path.dirname(input_path) or "."
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", "--outdir", out_dir, input_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120
        )
        converted = os.path.splitext(input_path)[0] + ".docx"
        if not os.path.exists(converted):
            raise RuntimeError("Conversion failed: output not found.")
        return converted
    except Exception as e:
        raise RuntimeError(f".doc -> .docx conversion error: {e}")

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    return "\n".join(lines)

def extract_text_from_pdf(path: str) -> str:
    parts = []
    with fitz.open(path) as pdf:
        for page in pdf:
            txt = page.get_text("text") or ""
            if txt.strip():
                parts.append(txt)
    return "\n".join(parts)

# -------------------------
# normalize input text (light cleaning)
# -------------------------
def normalize_text_for_llm(raw: str) -> str:
    if not raw:
        return ""
    lines = [ln.rstrip() for ln in raw.splitlines()]
    # drop repeated short header/footer lines
    freq = {}
    for ln in lines:
        k = ln.strip()
        if not k:
            continue
        freq[k] = freq.get(k, 0) + 1
    repeated = {ln for ln, c in freq.items() if c > 2 and len(ln) < 60}
    if repeated:
        lines = [ln for ln in lines if ln.strip() not in repeated]
    text = "\n".join(lines)
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

# -------------------------
# JSON extraction helpers (robust bracket balancing)
# -------------------------
def _find_balanced_block(text: str, open_ch: str, close_ch: str) -> Optional[str]:
    starts = [m.start() for m in re.finditer(re.escape(open_ch), text)]
    for start in starts:
        depth = 0
        for i in range(start, len(text)):
            ch = text[i]
            if ch == open_ch:
                depth += 1
            elif ch == close_ch:
                depth -= 1
                if depth == 0:
                    candidate = text[start : i + 1]
                    try:
                        json.loads(candidate)
                        return candidate
                    except Exception:
                        break
    return None

def first_json_candidate_from_text(text: str, prefer: str = "object") -> Optional[str]:
    if not text:
        return None
    cleaned = text.replace("```json", "").replace("```", "").strip()
    order = ["object", "array"] if prefer == "object" else ["array", "object"]
    for kind in order:
        blk = _find_balanced_block(cleaned, "{" , "}") if kind=="object" else _find_balanced_block(cleaned, "[", "]")
        if blk is not None:
            return blk
    return None

def parse_json_block(raw: str, expect: str = "object") -> Any:
    prefer = "object" if expect == "object" else "array"
    block = first_json_candidate_from_text(raw, prefer=prefer)
    if not block:
        return {} if expect == "object" else []
    try:
        data = json.loads(block)
    except Exception:
        return {} if expect == "object" else []
    if expect == "object":
        return data if isinstance(data, dict) else {}
    else:
        return data if isinstance(data, list) else []

# -------------------------
# Normalizers into universal schema
# -------------------------
def normalize_candidate_obj(obj: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "full_name": obj.get("full_name", "") or "",
        "email": obj.get("email", "") or "",
        "phone": obj.get("phone", "") or "",
        "address": obj.get("address", "") or "",
        "linkedin": obj.get("linkedin", "") or "",
        "portfolio": obj.get("portfolio", "") or "",
        "github": obj.get("github", "") or "",
        "other_profiles": obj.get("other_profiles", []) or [],
    }

def normalize_experience_list(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out = []
    for it in items or []:
        out.append({
            "company": it.get("company","") or "",
            "position": it.get("position","") or "",
            "start_date": it.get("start_date","") or "",
            "end_date": it.get("end_date","") or "",
            "location": it.get("location","") or "",
            "employment_type": it.get("employment_type","") or "",
            "environment": it.get("environment",[]) or [],
            "summary": it.get("summary",[]) or []
        })
    return out

def normalize_skills_obj(skills_list):
    """
    Convert:
      [{ "key": "...", "value": "..." }, ...]
    Into:
      { "key": "value", ... }
    """
    if not isinstance(skills_list, list):
        return {}

    normalized = {}
    for item in skills_list:
        key = item.get("key", "").strip()
        value = item.get("value", "").strip()
        if key:
            normalized[key] = value
    return normalized

def normalize_education_list(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out = []
    for it in items or []:
        out.append({
            "degree": it.get("degree","") or "",
            "institution": it.get("institution","") or "",
            "location": it.get("location","") or "",
            "start_year": it.get("start_year","") or "",
            "end_year": it.get("end_year","") or it.get("year","") or "",
            "gpa": it.get("gpa","") or "",
            "coursework": it.get("coursework",[]) or [],
            "awarded": it.get("awarded","") or ""
        })
    return out

def normalize_projects_list(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out = []
    for it in items or []:
        out.append({
            "title": it.get("title","") or "",
            "subtitle": it.get("subtitle","") or "",
            "organization": it.get("organization","") or "",
            "location": it.get("location","") or "",
            "start_date": it.get("start_date","") or "",
            "end_date": it.get("end_date","") or "",
            "technologies": it.get("technologies",[]) or [],
            "description": it.get("description",[]) or []
        })
    return out

TEMPLATE = """You are a resume data extraction engine.
Your task is to read the provided resume text and convert it into a structured JSON object that matches the schema below.
Do not shorten content. Do not merge bullet points. Preserve original meaning.

Return ONLY valid JSON. No explanations.

-------------------------
JSON SCHEMA
-------------------------

{{
  "candidate": {{
    "full_name": "",
    "email": "",
    "phone": "",
    "address": "",
    "linkedin": "",
    "github": "",
    "portfolio": ""
  }},

  "summary": [
    "bullet point 1",
    "bullet point 2",
    "bullet point 3"
  ],

  "skills": [
    {{
      "key": "Skill Category",
      "value": "Comma separated list of skills"
    }}
  ],

  "education": [
    {{
      "degree": "",
      "institution": "",
      "location": "",
      "year": "",
      "details": ""
    }}
  ],

  "experience": [
    {{
      "company": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "position": "",
      "summary": [
        "bullet point 1",
        "bullet point 2"
      ],
      "environment": "comma separated tools and technologies"
    }}
  ],

  "projects": [
    {{
      "title": "",
      "subtitle": "",
      "organization": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "technologies": "",
      "description": [
        "bullet point 1",
        "bullet point 2"
      ]
    }}
  ],

  "certifications": [
    {{
      "name": "",
      "issuer": "",
      "date": ""
    }}
  ]
}}

-------------------------
RULES
-------------------------

• All summary and experience items must be bullet-point arrays.
• Preserve dates exactly as written.
• Environment should be one comma-separated string.
• Each skill category becomes one row.
• Do not infer missing info; leave empty strings.
• Preserve the resume’s wording.
• Output must be valid JSON only.
• provide all resume information.


-------------------------
INPUT RESUME
-------------------------

{resumeText}

-------------------------
OUTPUT
-------------------------
Return JSON only.
"""

# -------------------------
# Cloud call logic with chunking for long resumes
# -------------------------
def chunk_text(text: str, chunk_size_chars: int = 4000) -> List[str]:
    """Split text into readable chunks (preserving line breaks where possible)."""
    if not text:
        return []
    parts = []
    start = 0
    L = len(text)
    while start < L:
        end = min(start + chunk_size_chars, L)
        # try to break on newline for nicer chunk boundaries
        if end < L:
            nl = text.rfind("\n", start, end)
            if nl > start:
                end = nl
        parts.append(text[start:end].strip())
        start = end
    return parts

def build_unified_prompt_from_text(text: str) -> str:
    chunks = chunk_text(text, chunk_size_chars=4200)
    if len(chunks) == 0:
        resume_chunks = "\"\""
    elif len(chunks) == 1:
        resume_chunks = f"\"\"\"{chunks[0]}\"\"\""
    else:
        # number the chunks so the model can reason across them
        numbered = []
        for i,c in enumerate(chunks, start=1):
            numbered.append(f"--- CHUNK {i}/{len(chunks)} START ---\n{c}\n--- CHUNK {i}/{len(chunks)} END ---")
        resume_chunks = "\"\"\"" + "\n\n".join(numbered) + "\"\"\""
    return TEMPLATE.format(resumeText=resume_chunks)
    # UNIVERSAL_PROMPT.format(resume_chunks=resume_chunks)

def extract_json_block(text: str):
    """
    Extracts the first valid JSON object { ... } from Ollama output.
    This handles reasoning text, comments, or anything after the JSON.
    """
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("No JSON object found in Ollama response.")
    json_str = match.group(0)
    return json.loads(json_str)



def run_cloud_unified_extraction(resume_text: str, timeout: int = 300):
    client = Client(
        host="https://ollama.com",
        headers={'Authorization': 'Bearer bce08f66335b43ed8feeca10d0b64876.g7V-_vkd4xJxkpbfWgG5YJKg'},
    )

    prompt = build_unified_prompt_from_text(resume_text)
    print("PROMPT:", prompt)
    messages = [{"role": "user", "content": prompt}]

    try:
        resp = client.chat(
            # model= "gpt-oss:120b-cloud",
            model= "gpt-oss:120b-cloud",
            messages=messages,
            stream=False
        )

        # print("RAW RESPONSE FROM OLLAMA:", resp)

        # 1. Pull raw content
        if hasattr(resp, "message") and hasattr(resp.message, "content"):
            content = resp.message.content
        elif isinstance(resp, dict) and "message" in resp and "content" in resp["message"]:
            content = resp["message"]["content"]
        else:
            content = str(resp)

    
        # 2. Normalize JSON (THE FIX)
        try:
            
            # print("TYPE:", type(content))
            parsed = json.loads(content)
            print("CONTENT FROM OLLAMA:", parsed)
            print("PARSED:", type(parsed))
            return parsed, {
                    "model": resp.get("model"),
                    "host": "https://ollama.com"
                }
        except Exception as e:
            return { "error": f"JSON parse error: {str(e)}" }, {
                "model": "gpt-oss:120b-cloud",
                "host": "https://ollama.com"
            }

        return parsed, {
            "model": getattr(resp, "model", "gpt-oss:120b-cloud"),
            "host": "https://ollama.com"
        }

    except Exception as e:
        return { "error": f"Ollama cloud call error: {e}" }, {
            "model": "gpt-oss:120b-cloud",
            "host": "https://ollama.com"
        }

# -------------------------
# Extraction endpoint and wiring
# -------------------------
def extract_json_from_text_unified(resume_text: str, src_filename: str) -> Tuple[Dict[str,Any], Dict[str,str]]:
    """
    One-call unified extraction. Returns (final_json_schema, raw_outs)
    """
    raw_out = {}
    prompt_out, meta = run_cloud_unified_extraction(resume_text)
    raw_out["meta"] = meta
    parsed = prompt_out
    if not parsed or not isinstance(parsed, dict):
        final = {
            "source_file": src_filename,
            "metadata": {"schema_reference": "Universal Resume Schema v1", "extracted_from": src_filename},
            "candidate": normalize_candidate_obj({}),
            "professional_summary": "",
            "skills": normalize_skills_obj({}),
            "employment_history": [],
            "education": [],
            "projects": []
        }
        raw_out["parse_error"] = "Could not parse JSON from model output. See unified_raw."
        return final, raw_out

    candidate_norm = normalize_candidate_obj(parsed.get("candidate", {})) or {}
    prof_summary = parsed.get("summary", []) or []
    skills_norm = normalize_skills_obj(parsed.get("skills", []))
    exp_norm = normalize_experience_list(parsed.get("experience", [])) or []
    edu_norm = normalize_education_list(parsed.get("education", [])) or []
    proj_norm = normalize_projects_list(parsed.get("projects", [])) or []
    cert_norm = parsed.get("certifications", [])
    final = {
        "source_file": src_filename,
        "metadata": {"schema_reference": "Universal Resume Schema v1", "extracted_from": src_filename},
        "candidate": candidate_norm,
        "professional_summary": prof_summary,
        "skills": skills_norm,
        "employment_history": exp_norm,
        "education": edu_norm,
        "projects": proj_norm,
        "certifications": cert_norm
    }
    return final, raw_out

@app.post("/extract_json")
async def extract_json(resume: UploadFile = File(...), use_local_path: Optional[str] = Form(None)):
    """
    Upload a resume file or provide a local path (dev only) to extract JSON using Ollama Cloud unified prompt.
    - resume: file upload (PDF/DOCX/DOC)
    - use_local_path: optional form string; when present, the server will read that file path instead of the upload (dev/testing)
    """
    try:
        # read file either from upload or local path
        if use_local_path:
            if not os.path.exists(use_local_path):
                return JSONResponse(status_code=404, content={"error": f"Local path not found: {use_local_path}"})
            saved_path = use_local_path
            filename = os.path.basename(saved_path)
        else:
            filename = resume.filename or "uploaded_resume"
            ext = os.path.splitext(filename)[1].lower()
            if ext not in [".pdf", ".docx", ".doc"]:
                return JSONResponse(status_code=400, content={"error": "Unsupported file type. Use PDF, DOCX, or DOC."})
            saved_path = os.path.join(OUTPUT_DIR, filename)
            with open(saved_path, "wb") as f:
                f.write(await resume.read())
            if ext == ".doc":
                saved_path = convert_doc_to_docx(saved_path)

        # extract raw text
        if saved_path.lower().endswith(".docx"):
            raw_text = extract_text_from_docx(saved_path)
        else:
            raw_text = extract_text_from_pdf(saved_path)

        normalized = normalize_text_for_llm(raw_text)

        final_struct, raw_debug = extract_json_from_text_unified(normalized, filename)

        # Save output files for debugging
        out_json_path = os.path.join(OUTPUT_DIR, os.path.splitext(filename)[0] + "_extracted.json")
        with open(out_json_path, "w", encoding="utf-8") as jf:
            json.dump(final_struct, jf, indent=2)

        debug_path = os.path.join(OUTPUT_DIR, os.path.splitext(filename)[0] + "_sections_debug.json")
        with open(debug_path, "w", encoding="utf-8") as df:
            json.dump(raw_debug, df, indent=2)

        return JSONResponse(content=final_struct, status_code=200)

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

# -------------------------
# Remaining supporting endpoints (preview_html, debug, templates listing)
# -------------------------
def generate_html_from_schema(data: Dict[str, Any]) -> str:
    c = data.get("candidate", {})
    exp = data.get("employment_history", [])
    edu = data.get("education", [])
    skills = data.get("skills", {})
    topline = skills.get("topline", []) or []
    tech = skills.get("technical", {}) or {}
    summary = data.get("professional_summary", "")

    def esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    html = f"""<!doctype html><html lang="en"><head><meta charset="utf-8" /><title>{esc(c.get("full_name","Resume Preview"))}</title>
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <style>body{{font-family:Arial,Helvetica,sans-serif;background:#f7f8fa;margin:0;padding:20px}}.container{{max-width:960px;margin:0 auto;background:#fff;padding:20px;border-radius:8px}}</style>
    </head><body><div class="container">"""
    html += f"<h1>{esc(c.get('full_name',''))}</h1>"
    html += f"<p><strong>Email:</strong> {esc(c.get('email',''))} &nbsp; <strong>Phone:</strong> {esc(c.get('phone',''))}</p>"
    html += f"<h3>Summary</h3><p>{esc(summary) or 'No summary parsed.'}</p>"
    html += "<h3>Skills</h3>"
    for t in topline:
        html += f"<span style='display:inline-block;padding:4px 8px;margin:4px;background:#eef2ff;border-radius:999px;font-size:12px'>{esc(t)}</span>"
    for k, arr in tech.items():
        if arr:
            html += f"<p><strong>{esc(k.capitalize())}:</strong> {', '.join(esc(x) for x in arr)}</p>"
    html += "<h3>Experience</h3>"
    for j in exp:
        html += f"<div style='margin-bottom:14px'><strong>{esc(j.get('role',''))}</strong> — {esc(j.get('company',''))}<br><small>{esc(j.get('start_date',''))} – {esc(j.get('end_date',''))} | {esc(j.get('location',''))}</small>"
        if j.get("bullets"):
            html += "<ul>"
            for b in j.get("bullets",[]):
                html += f"<li>{esc(b)}</li>"
            html += "</ul>"
        html += "</div>"
    html += "<h3>Education</h3>"
    for e in edu:
        html += f"<p><strong>{esc(e.get('degree',''))}</strong> — {esc(e.get('institution',''))} ({esc(e.get('end_year',''))})</p>"
    html += "</div></body></html>"
    return html

@app.post("/preview_html")
async def preview_html(resume: UploadFile = File(...)):
    extract_resp = await extract_json(resume)
    if isinstance(extract_resp, JSONResponse) and extract_resp.status_code != 200:
        return extract_resp
    data = extract_resp.body
    if isinstance(data, (bytes, bytearray)):
        data = json.loads(data.decode("utf-8"))
    filename = resume.filename or "uploaded_resume"
    html = generate_html_from_schema(data)
    html_file = os.path.join(OUTPUT_DIR, os.path.splitext(filename)[0] + "_preview.html")
    with open(html_file, "w", encoding="utf-8") as hf:
        hf.write(html)
    return JSONResponse(content={"message": "HTML preview generated.", "html_file": os.path.basename(html_file), "download_html": f"/download/{os.path.basename(html_file)}"}, status_code=200)

@app.post("/debug")
async def debug_endpoint(resume: UploadFile = File(...)):
    """
    Returns normalized text preview and raw model output to help debugging parse errors.
    """
    try:
        filename = resume.filename or "uploaded_resume"
        ext = os.path.splitext(filename)[1].lower()
        saved_path = os.path.join(OUTPUT_DIR, filename)
        with open(saved_path, "wb") as f:
            f.write(await resume.read())
        if ext == ".doc":
            saved_path = convert_doc_to_docx(saved_path)
        raw_text = extract_text_from_docx(saved_path) if saved_path.lower().endswith(".docx") else extract_text_from_pdf(saved_path)
        normalized = normalize_text_for_llm(raw_text)
        final_struct, raw_out = extract_json_from_text_unified(normalized, filename)
        return JSONResponse(content={"normalized_text_preview": normalized[:3000], "structured_preview": final_struct, "raw_out": raw_out}, status_code=200)
    except Exception as e:
        return JSONResponse(status_code=500, content={"fatal_debug_error": str(e)})

@app.get("/templates")
def list_templates():
    return {"templates": sorted([f for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith(".docx") or f.lower().endswith(".doc")])}

# -------------------------
# Template formatting code (same style as your prior logic)
# -------------------------
EDU_HEADERS = {"degree","area","area of study","major","program","school","institution","location","awarded","date","year"}
EMP_HEADERS = {"role","title","position","company","employer","organization","start","end","from","to","dates","location","duty","responsibility","bullets"}
SKILL_HEADERS = {"skill","category","details","skill name","skill / category","skill/category","skill/ category","skill/details"}
CERT_HEADERS = {"certification","certificate","authority","issued","license"}

def normalize_cell_text(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())

def table_header_tokens(table: Table) -> List[str]:
    if not table.rows:
        return []
    header_row = table.rows[0]
    tokens = []
    for cell in header_row.cells:
        txt = normalize_cell_text(cell.text)
        tokens.append(txt)
    return tokens

def detect_table_type(table: Table) -> Optional[str]:
    tokens = table_header_tokens(table)
    words = set()
    for t in tokens:
        if not t:
            continue
        for w in re.split(r"[\/\|\-,\s]+", t):
            words.add(w.strip())
    if words & EDU_HEADERS:
        return "education"
    if words & EMP_HEADERS:
        return "employment"
    if words & SKILL_HEADERS:
        return "skills"
    if words & CERT_HEADERS:
        return "certifications"
    joined = " ".join(tokens)
    if any(k in joined for k in ["degree","institution","school","major"]):
        return "education"
    if any(k in joined for k in ["company","role","employer","position","title"]):
        return "employment"
    if any(k in joined for k in ["skill","category"]):
        return "skills"
    return None

# -------------------------
# Robust paragraph helpers (single authoritative implementation)
# -------------------------
def _remove_paragraph(paragraph):
    """
    Remove a paragraph element from the document.
    """
    if paragraph is None:
        return
    elem = getattr(paragraph, "_element", None)
    if elem is None:
        return
    parent = elem.getparent()
    if parent is None:
        return
    try:
        parent.remove(elem)
    except Exception:
        logger.exception("Failed to remove paragraph element")
    # detach python-docx wrapper
    try:
        paragraph._p = paragraph._element = None
    except Exception:
        pass

def _ensure_paragraph_has_run(paragraph):
    """
    Ensure the paragraph has at least one run so we can safely write text into it.
    Returns the first run.
    """
    if paragraph is None:
        raise ValueError("paragraph is None")
    runs = list(getattr(paragraph, "runs", []))
    if runs:
        return runs[0]
    # create a run by adding empty text
    try:
        r = paragraph.add_run("")
        return r
    except Exception:
        # As a last resort, add a w:r element to the underlying XML
        try:
            r_elm = OxmlElement("w:r")
            t_elm = OxmlElement("w:t")
            t_elm.text = ""
            r_elm.append(t_elm)
            paragraph._element.append(r_elm)
            runs = list(getattr(paragraph, "runs", []))
            if runs:
                return runs[-1]
        except Exception:
            logger.exception("Failed to ensure run on paragraph")
    raise RuntimeError("Unable to create a run in paragraph")

def _safe_add_paragraph_after(doc_obj, ref_paragraph, text="", style=None):
    """
    Insert a new paragraph immediately after ref_paragraph, if possible.
    If insertion after is not possible (detached paragraph or unsupported parent),
    fall back to appending to the ref_paragraph's parent via its add_paragraph method,
    or finally append to the document object.
    Returns a python-docx Paragraph object with text already added (if provided).
    """
    # Quick helpers
    def _create_paragraph_via_add_next():
        # Create an empty <w:p> element and wrap it
        new_p_elm = OxmlElement("w:p")
        ref_paragraph._element.addnext(new_p_elm)
        return Paragraph(new_p_elm, ref_paragraph._parent)

    def _create_paragraph_via_parent_add():
        parent = getattr(ref_paragraph, "_parent", None)
        if parent is not None and hasattr(parent, "add_paragraph"):
            return parent.add_paragraph(text)
        return None

    # 1) Try the ideal path: insert after the existing paragraph element
    try:
        if getattr(ref_paragraph, "_element", None) is not None and ref_paragraph._element.getparent() is not None and getattr(ref_paragraph, "_parent", None) is not None:
            new_para = _create_paragraph_via_add_next()
            # Ensure there's at least one run to write into
            try:
                run = _ensure_paragraph_has_run(new_para)
                if text:
                    # set text into the first run (safer than add_run on some detached wrappers)
                    try:
                        run.text = str(text)
                    except Exception:
                        # fallback to add_run if run.text assignment fails
                        new_para.add_run(str(text))
                if style:
                    try:
                        new_para.style = style
                    except Exception:
                        logger.debug("Could not set style %s on new paragraph", style, exc_info=True)
                return new_para
            except Exception:
                logger.debug("Failed to ensure run on new paragraph created via addnext", exc_info=True)
    except Exception:
        logger.debug("Insertion via addnext failed, will try parent.add_paragraph fallback", exc_info=True)

    # 2) Try adding via the parent (appends to the parent's end)
    try:
        para_via_parent = _create_paragraph_via_parent_add()
        if para_via_parent is not None:
            # parent.add_paragraph already put text if we passed it; ensure style
            if style:
                try:
                    para_via_parent.style = style
                except Exception:
                    logger.debug("Could not set style %s on parent-added paragraph", style, exc_info=True)
            return para_via_parent
    except Exception:
        logger.debug("Parent.add_paragraph fallback failed", exc_info=True)

    # 3) Final fallback: append to the main document
    try:
        fallback = doc_obj.add_paragraph(text)
        if style:
            try:
                fallback.style = style
            except Exception:
                logger.debug("Could not set style %s on fallback paragraph", style, exc_info=True)
        return fallback
    except Exception:
        logger.exception("Final fallback to doc.add_paragraph failed; creating empty Paragraph wrapper")

    # 4) As last-resort, create an empty OXML paragraph and wrap with a guessed parent (document)
    try:
        new_elm = OxmlElement("w:p")
        parent_wrapper = getattr(doc_obj, "_body", None) or getattr(doc_obj, "_element", None)
        if parent_wrapper is not None:
            parent_wrapper.append(new_elm)
            return Paragraph(new_elm, getattr(doc_obj, "_element", doc_obj))
    except Exception:
        logger.exception("Absolute last-resort paragraph creation failed")

    # If everything fails, raise so caller can handle it
    raise RuntimeError("Unable to create a new paragraph in this document/context")

def _replace_paragraph_text(paragraph, new_text: str):
    """
    Replace the full text of a paragraph by removing all runs and adding a single run.
    """
    if paragraph is None:
        return
    # remove all run elements
    for run in list(getattr(paragraph, "runs", [])):
        try:
            paragraph._element.remove(run._element)
        except Exception:
            logger.debug("Failed to remove run element during replace", exc_info=True)
    # add single run
    paragraph.add_run(str(new_text))

# -------------------------
# Skills parsing helper (for free-text blobs) and formatting
# -------------------------
_SKILLS_DICT_LIKE_RE = _re.compile(r"^\s*(?:\{.*:.*\}|\[.*\])\s*$", _re.DOTALL)


def build_skills_text(skills_obj: Dict[str, Any]) -> str:
    """
    Build a multi-line text representation of skills.

    - If skills_obj contains the nested shape with 'topline'/'technical', keep that behavior.
    - Otherwise iterate over the keys dynamically (preserve order) and render each key -> comma-joined values.
    """
    lines = []

    if not skills_obj:
        return ""

    # If it's a list-of-pairs like [{"key":"...","value":"..."}], convert to a flat dict
    if isinstance(skills_obj, list):
        flat = {}
        for it in skills_obj:
            if not isinstance(it, dict):
                continue
            k = (it.get("key") or it.get("name") or "").strip()
            v = it.get("value", "") if "value" in it else it.get("values", "") or ""
            if k:
                flat[k] = v
        skills_obj = flat

    # Special handling for nested structure (preserve previous behavior)
    if isinstance(skills_obj, dict) and ("topline" in skills_obj or "technical" in skills_obj):
        topline = skills_obj.get("topline", []) or []
        if topline:
            lines.append("Topline: " + ", ".join(str(x).strip() for x in topline if x is not None and str(x).strip()))
        tech = skills_obj.get("technical", {}) or {}
        # iterate technical keys dynamically as well (preserve insertion order)
        for tech_key, tech_val in tech.items():
            if not tech_val:
                continue
            # normalize tech_val to comma-joined string
            if isinstance(tech_val, (list, tuple)):
                vals = ", ".join(str(x).strip() for x in tech_val if x is not None and str(x).strip())
            else:
                sval = str(tech_val).strip()
                # try json decode if it looks like a list
                try:
                    maybe = json.loads(sval)
                    if isinstance(maybe, (list, tuple)):
                        vals = ", ".join(str(x).strip() for x in maybe if x is not None and str(x).strip())
                    else:
                        vals = re.sub(r"\s+", " ", sval)
                except Exception:
                    vals = re.sub(r"\s+", " ", sval)
            if vals:
                # Capitalize the tech_key label nicely
                lines.append(f"{tech_key.capitalize()}: {vals}")
        if skills_obj.get("soft_skills"):
            ss = skills_obj.get("soft_skills", [])
            if isinstance(ss, (list, tuple)):
                lines.append("Soft Skills: " + ", ".join(str(x).strip() for x in ss if x is not None and str(x).strip()))
            else:
                s = str(ss).strip()
                if s:
                    lines.append("Soft Skills: " + re.sub(r"\s+", " ", s))
        if skills_obj.get("certifications"):
            cert = skills_obj.get("certifications", [])
            if isinstance(cert, (list, tuple)):
                lines.append("Certifications: " + ", ".join(str(x).strip() for x in cert if x is not None and str(x).strip()))
            else:
                s = str(cert).strip()
                if s:
                    lines.append("Certifications: " + re.sub(r"\s+", " ", s))
        return "\n".join(lines)

    # Otherwise treat as flat mapping category -> (string or list), iterating keys dynamically
    if isinstance(skills_obj, dict):
        for k, v in skills_obj.items():
            if v is None:
                continue
            # normalize key label (preserve original case, but trim)
            label = str(k).strip()
            if not label:
                continue

            vals = ""
            # value is a list/tuple
            if isinstance(v, (list, tuple)):
                vals = ", ".join(str(x).strip() for x in v if x is not None and str(x).strip())
            else:
                s = str(v).strip()
                if not s:
                    continue
                # try json decode if the value is a JSON list encoded as string
                try:
                    maybe = json.loads(s)
                    if isinstance(maybe, (list, tuple)):
                        vals = ", ".join(str(x).strip() for x in maybe if x is not None and str(x).strip())
                    else:
                        vals = re.sub(r"\s+", " ", s)
                except Exception:
                    # not JSON: treat comma separated or raw string
                    vals = re.sub(r"\s+", " ", s)

            if vals:
                lines.append(f"{label}: {vals}")

        return "\n".join(lines)

    # Fallback: stringify whatever it is
    try:
        return str(skills_obj)
    except Exception:
        return ""

# -------------------------
# Paragraph processing and placeholder replacement
# -------------------------

def add_formatted_skill_line(paragraph, key: str, value: str):
    """
    Adds formatted skill line to a paragraph:
    - key bold
    - Calibri font
    - font size 11
    """
    # Ensure paragraph exists
    if paragraph is None:
        return
    # clear any existing runs in the paragraph to avoid mixing styles unexpectedly
    try:
        for r in list(getattr(paragraph, "runs", [])):
            try:
                paragraph._element.remove(r._element)
            except Exception:
                pass
    except Exception:
        pass

    run_key = paragraph.add_run(f"{key}: ")
    try:
        run_key.bold = True
        run_key.font.name = "Calibri"
        run_key.font.size = Pt(11)
    except Exception:
        # ignore styling failures (font may not be available)
        pass

    run_val = paragraph.add_run(value or "")
    try:
        run_val.bold = False
        run_val.font.name = "Calibri"
        run_val.font.size = Pt(11)
    except Exception:
        pass

def insert_skills_table(doc, ref_paragraph, skills_obj: Dict[str, Any]):
    if ref_paragraph is None or getattr(ref_paragraph, "_element", None) is None:
        return doc.add_paragraph("")

    if isinstance(skills_obj, list):
        flat_list = []
        for it in skills_obj:
            if isinstance(it, dict):
                if "value" in it and it.get("value"):
                    v = it.get("value")
                    if isinstance(v, (list, tuple)):
                        flat_list.extend([str(x).strip() for x in v if str(x).strip()])
                    else:
                        flat_list.extend([x.strip() for x in re.split(r",\s*", str(v)) if x.strip()])
                else:
                    k = it.get("key") or it.get("name") or ""
                    if k:
                        flat_list.append(k.strip())
            else:
                s = str(it).strip()
                if s:
                    flat_list.append(s)
        skills_list = flat_list
    elif isinstance(skills_obj, dict):
        # if OCFO_MODE prefer flattening values into bullets
        skills_list = []
        for k, v in skills_obj.items():
            if isinstance(v, (list, tuple)):
                skills_list.extend([str(x).strip() for x in v if str(x).strip()])
            else:
                s = str(v).strip()
                if s:
                    try:
                        maybe = json.loads(s)
                        if isinstance(maybe, (list, tuple)):
                            skills_list.extend([str(x).strip() for x in maybe if str(x).strip()])
                        else:
                            skills_list.extend([x.strip() for x in re.split(r",\s*", s) if x.strip()])
                    except Exception:
                        skills_list.extend([x.strip() for x in re.split(r",\s*", s) if x.strip()])
        if not skills_list:
            skills_list = [k for k in skills_obj.keys() if (k and str(k).strip())]
    else:
        s = str(skills_obj or "").strip()
        if not s:
            skills_list = []
        else:
            try:
                parsed = json.loads(s)
                if isinstance(parsed, (list, tuple)):
                    skills_list = [str(x).strip() for x in parsed if str(x).strip()]
                elif isinstance(parsed, dict):
                    skills_list = [str(x).strip() for x in parsed.keys() if str(x).strip()]
                else:
                    skills_list = [x.strip() for x in re.split(r",\s*", s) if x.strip()]
            except Exception:
                skills_list = [x.strip() for x in re.split(r",\s*", s) if x.strip()]

    def _zero_para_spacing(p):
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        except Exception:
            pass

    if OCFO_MODE:
        insert_ref = ref_paragraph
        before, sep, after = (ref_paragraph.text or "").partition("{{SKILLS}}") if "{{SKILLS}}" in (ref_paragraph.text or "") else (None, None, None)
        if before and before.strip():
            insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=before)
        for item in skills_list:
            p = _safe_add_paragraph_after(doc, insert_ref, text="")
            _zero_para_spacing(p)
            try:
                p.style = "List Bullet"
            except Exception:
                pass
            try:
                run = p.add_run(str(item))
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            except Exception:
                p.add_run(str(item))
            insert_ref = p
        if after and after.strip():
            _safe_add_paragraph_after(doc, insert_ref, text=after)
        try:
            _remove_paragraph(ref_paragraph)
        except Exception:
            pass
        spacer = _safe_add_paragraph_after(doc, insert_ref, text="")
        _zero_para_spacing(spacer)
        return spacer

    # fall back to original table behavior if not OCFO_MODE
    # (recreate the previous multi-key / single-key logic)
    if not skills_list:
        new_p = OxmlElement("w:p")
        ref_paragraph._element.addnext(new_p)
        return Paragraph(new_p, getattr(ref_paragraph, "_parent", doc))

    if len(skills_list) <= 12 and len(skills_list) > 0:
        ncols = 4
        rows = (len(skills_list) + ncols - 1) // ncols
        tbl_elm = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr")
        tbl_elm.append(tbl_pr)
        tbl_grid = OxmlElement("w:tblGrid")
        for _ in range(ncols):
            g = OxmlElement("w:gridCol")
            g.set(qn("w:w"), "2400")
            tbl_grid.append(g)
        tbl_elm.append(tbl_grid)
        try:
            section = doc.sections[0]
        except Exception:
            section = None
        if section is not None:
            content_width = section.page_width - section.left_margin - section.right_margin
            twips_width = int(content_width / 635)
            tblPr = tbl_elm.find(qn("w:tblPr"))
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl_elm.insert(0, tblPr)
            tblW = tblPr.find(qn("w:tblW"))
            if tblW is None:
                tblW = OxmlElement("w:tblW")
                tblPr.append(tblW)
            tblW.set(qn("w:w"), str(twips_width))
            tblW.set(qn("w:type"), "dxa")
        ref_paragraph._element.addnext(tbl_elm)
        parent = getattr(ref_paragraph, "_parent", doc)
        table = _Table(tbl_elm, parent)
        for _ in range(rows):
            table.add_row()
        idx = 0
        for r in range(rows):
            for c in range(ncols):
                cell = table.cell(r, c)
                if idx < len(skills_list):
                    for p in list(cell.paragraphs):
                        try:
                            cell._element.remove(p._element)
                        except Exception:
                            pass
                    p = cell.add_paragraph()
                    _zero_para_spacing(p)
                    try:
                        p.style = "List Bullet"
                        rr = p.add_run(skills_list[idx])
                        rr.font.name = "Calibri"
                        rr.font.size = Pt(11)
                    except Exception:
                        p.add_run(skills_list[idx])
                else:
                    p = cell.add_paragraph()
                    _zero_para_spacing(p)
                idx += 1
        try:
            tbl = table._element
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{border_name}')
                b.set(qn('w:val'), 'nil')
                tblBorders.append(b)
            tblPr.append(tblBorders)
        except Exception:
            pass
        new_p = OxmlElement("w:p")
        table._element.addnext(new_p)
        return Paragraph(new_p, parent)

    new_p = OxmlElement("w:p")
    ref_paragraph._element.addnext(new_p)
    return Paragraph(new_p, getattr(ref_paragraph, "_parent", doc))

# -------------------------
# New: Education / Certifications / Projects insertion helpers
# -------------------------
# ---------- Section-placement helpers ----------

def _get_paragraph_index(doc, paragraph):
    try:
        for i, p in enumerate(list(doc.paragraphs)):
            if p is paragraph:
                return i
    except Exception:
        pass
    try:
        for i, p in enumerate(list(doc.paragraphs)):
            if p.text == paragraph.text:
                return i
    except Exception:
        pass
    return -1

def _looks_like_employment_placeholder(key: str) -> bool:
    if not key:
        return False
    k = key.strip().upper()
    if "EMPLOY" in k or "EMPLOYMENT" in k or "WORK" in k or "EXPERIENCE" in k:
        return True
    return False

def _list_looks_like_experience_items(lst: List[Any]) -> bool:
    if not isinstance(lst, list) or not lst:
        return False
    expected_keys = {"company", "position", "title", "role", "start_date", "end_date", "summary", "bullets"}
    for it in lst:
        if isinstance(it, dict):
            if any(k in it for k in expected_keys):
                return True
    return False

def _zero_para_spacing(paragraph):
    try:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
    except Exception:
        pass

def insert_experience_block(doc, ref_paragraph, experiences: List[Dict[str,Any]]):
    if ref_paragraph is None or getattr(ref_paragraph, "_element", None) is None:
        ref_paragraph = doc.add_paragraph("")
    insert_ref = ref_paragraph
    for entry in experiences or []:
        company = (entry.get("company") or "").strip()
        location = (entry.get("location") or "").strip()
        position = (entry.get("position") or entry.get("title") or entry.get("role") or "").strip()
        start = (entry.get("start_date") or "").strip()
        end = (entry.get("end_date") or "").strip()
        bullets = entry.get("summary") or entry.get("bullets") or entry.get("details") or []
        env_raw = entry.get("environment") or entry.get("environment_tools") or entry.get("environment_list") or ""
        tbl_elm = OxmlElement("w:tbl")
        tbl_pr = OxmlElement("w:tblPr"); tbl_elm.append(tbl_pr)
        tbl_grid = OxmlElement("w:tblGrid")
        c1 = OxmlElement("w:gridCol"); c1.set(qn("w:w"), str(6000)); tbl_grid.append(c1)
        c2 = OxmlElement("w:gridCol"); c2.set(qn("w:w"), str(3000)); tbl_grid.append(c2)
        tbl_elm.append(tbl_grid)
        try:
            insert_ref._element.addnext(tbl_elm)
            parent = getattr(insert_ref, "_parent", None) or getattr(doc, "_element", None)
            table = _Table(tbl_elm, parent)
        except Exception:
            table = doc.add_table(rows=0, cols=2)
        try:
            tbl_root = table._tbl
            for tr in list(tbl_root.findall(qn('w:tr'))):
                tbl_root.remove(tr)
        except Exception:
            pass
        row = table.add_row()
        left_cell = row.cells[0]; right_cell = row.cells[1]
        for p in list(left_cell.paragraphs):
            try:
                left_cell._element.remove(p._element)
            except Exception:
                pass
        for p in list(right_cell.paragraphs):
            try:
                right_cell._element.remove(p._element)
            except Exception:
                pass
        p_left = left_cell.add_paragraph()
        _zero_para_spacing(p_left)
        comp_text = company + (", " + location if location else "")
        run_left = p_left.add_run(comp_text)
        try:
            run_left.bold = True
            run_left.font.name = "Calibri"
            run_left.font.size = Pt(12)
        except Exception:
            pass
        p_right = right_cell.add_paragraph()
        _zero_para_spacing(p_right)
        try:
            p_right.alignment = 2
        except Exception:
            pass
        date_text = (start + " – " + end).strip(" –")
        rdate = p_right.add_run(date_text)
        try:
            rdate.bold = True
            rdate.font.name = "Calibri"
            rdate.font.size = Pt(11)
        except Exception:
            pass
        try:
            new_p = OxmlElement("w:p")
            table._element.addnext(new_p)
            insert_ref = Paragraph(new_p, getattr(insert_ref, "_parent", doc))
        except Exception:
            insert_ref = doc.add_paragraph("")
        if position:
            p_role = _safe_add_paragraph_after(doc, insert_ref, text="")
            _zero_para_spacing(p_role)
            r_role = p_role.add_run(position)
            try:
                r_role.bold = True
                r_role.font.name = "Calibri"
                r_role.font.size = Pt(11)
            except Exception:
                pass
            insert_ref = p_role
        if isinstance(bullets, str):
            bls = [ln.strip() for ln in bullets.splitlines() if ln.strip()] if "\n" in bullets else ([bullets.strip()] if bullets.strip() else [])
        elif isinstance(bullets, (list, tuple)):
            bls = [str(x).strip() for x in bullets if x is not None and str(x).strip()]
        else:
            bls = []
        for b in bls:
            p_b = _safe_add_paragraph_after(doc, insert_ref, text="")
            _zero_para_spacing(p_b)
            try:
                p_b.style = "List Bullet"
                run_b = p_b.add_run(str(b))
                try:
                    run_b.font.name = "Calibri"
                    run_b.font.size = Pt(11)
                except Exception:
                    pass
            except Exception:
                run_b = p_b.add_run("\u2022 " + str(b))
                try:
                    run_b.font.name = "Calibri"
                    run_b.font.size = Pt(11)
                except Exception:
                    pass
            insert_ref = p_b
        env_items = []
        if isinstance(env_raw, (list, tuple)):
            env_items = [str(x).strip() for x in env_raw if x is not None and str(x).strip()]
        else:
            s = str(env_raw or "").strip()
            if s:
                try:
                    maybe = json.loads(s)
                    if isinstance(maybe, (list, tuple)):
                        env_items = [str(x).strip() for x in maybe if x is not None and str(x).strip()]
                    else:
                        env_items = [x.strip() for x in re.split(r",\s*", s) if x.strip()]
                except Exception:
                    env_items = [x.strip() for x in re.split(r",\s*", s) if x.strip()]
        if env_items:
            p_env = _safe_add_paragraph_after(doc, insert_ref, text="")
            _zero_para_spacing(p_env)
            run_label = p_env.add_run("Environment: ")
            try:
                run_label.bold = True
                run_label.font.name = "Calibri"
                run_label.font.size = Pt(11)
            except Exception:
                pass
            run_vals = p_env.add_run(", ".join(env_items))
            try:
                run_vals.bold = False
                run_vals.font.name = "Calibri"
                run_vals.font.size = Pt(11)
            except Exception:
                pass
            insert_ref = p_env
        spacer = _safe_add_paragraph_after(doc, insert_ref, text="")
        _zero_para_spacing(spacer)
        insert_ref = spacer
    return insert_ref

_SECTION_HEADINGS = [
    "summary", "skills", "skill", "experience", "employment", "work", "education",
    "certifications", "certifications and licenses", "projects"
]

def find_section_end_paragraph(doc, heading_paragraph):
    if heading_paragraph is None:
        return None
    paras = list(doc.paragraphs)
    start_idx = _get_paragraph_index(doc, heading_paragraph)
    if start_idx < 0:
        return heading_paragraph
    def looks_like_heading_text(txt: str):
        if not txt:
            return False
        s = txt.strip().lower()
        for h in _SECTION_HEADINGS:
            if re.match(rf"^\s*{re.escape(h)}\b", s):
                return True
        return False
    i = start_idx + 1
    last_idx = start_idx
    while i < len(paras):
        txt = (paras[i].text or "").strip()
        if txt and looks_like_heading_text(txt):
            break
        last_idx = i
        i += 1
    return paras[last_idx] if last_idx >= 0 else heading_paragraph

def insert_education_block(doc, heading_paragraph, education_list: List[Dict[str,Any]]):
    end_para = find_section_end_paragraph(doc, heading_paragraph) or heading_paragraph
    ref = end_para
    p_heading = _safe_add_paragraph_after(doc, ref, text="")
    _zero_para_spacing(p_heading)
    try:
        run_h = p_heading.add_run("Education")
        run_h.bold = True
        run_h.font.name = "Calibri"
        run_h.font.size = Pt(11)
    except Exception:
        p_heading.add_run("Education")
    insert_ref = p_heading
    if not education_list:
        spacer = _safe_add_paragraph_after(doc, insert_ref, text="")
        _zero_para_spacing(spacer)
        return spacer
    tbl_elm = OxmlElement("w:tbl")
    tbl_pr = OxmlElement("w:tblPr"); tbl_elm.append(tbl_pr)
    tbl_grid = OxmlElement("w:tblGrid")
    col1 = OxmlElement("w:gridCol"); col1.set(qn("w:w"), str(8000))
    col2 = OxmlElement("w:gridCol"); col2.set(qn("w:w"), str(1200))
    tbl_grid.append(col1); tbl_grid.append(col2)
    tbl_elm.append(tbl_grid)
    try:
        ref._element.addnext(tbl_elm)
        parent = getattr(ref, "_parent", None) or getattr(doc, "_element", None)
        table = _Table(tbl_elm, parent)
    except Exception:
        table = doc.add_table(rows=0, cols=2)
    try:
        tbl_root = table._tbl
        for tr in list(tbl_root.findall(qn('w:tr'))):
            tbl_root.remove(tr)
    except Exception:
        pass
    for edu in education_list or []:
        degree = (edu.get("degree") or "").strip()
        inst = (edu.get("institution") or "").strip()
        year = (edu.get("end_year") or edu.get("year") or "").strip()
        left_text_parts = []
        if degree:
            left_text_parts.append(degree)
        if inst:
            left_text_parts.append(inst)
        left_text = " — ".join(left_text_parts) if left_text_parts else ""
        row = table.add_row()
        left_cell = row.cells[0]
        right_cell = row.cells[1]
        for p in list(left_cell.paragraphs):
            try:
                left_cell._element.remove(p._element)
            except Exception:
                pass
        for p in list(right_cell.paragraphs):
            try:
                right_cell._element.remove(p._element)
            except Exception:
                pass
        p_left = left_cell.add_paragraph()
        _zero_para_spacing(p_left)
        try:
            r_left = p_left.add_run(left_text)
            r_left.font.name = "Calibri"
            r_left.font.size = Pt(11)
            if degree and inst:
                try:
                    p_left._element.remove(r_left._element)
                except Exception:
                    pass
                run_deg = p_left.add_run(degree)
                try:
                    run_deg.bold = True
                    run_deg.font.name = "Calibri"
                    run_deg.font.size = Pt(11)
                except Exception:
                    pass
                run_sep = p_left.add_run(" — " + inst)
                try:
                    run_sep.font.name = "Calibri"
                    run_sep.font.size = Pt(11)
                except Exception:
                    pass
            else:
                pass
        except Exception:
            p_left.add_run(left_text)
        p_right = right_cell.add_paragraph()
        _zero_para_spacing(p_right)
        try:
            p_right.alignment = 2
        except Exception:
            pass
        try:
            r_right = p_right.add_run(year or "")
            r_right.font.name = "Calibri"
            r_right.font.size = Pt(11)
        except Exception:
            p_right.add_run(year or "")
    try:
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top','left','bottom','right','insideH','insideV'):
            border_elm = OxmlElement(f'w:{border_name}')
            border_elm.set(qn('w:val'), 'nil')
            tblBorders.append(border_elm)
        tblPr.append(tblBorders)
    except Exception:
        pass
    try:
        sect = doc.sections[0]
        page_width = sect.page_width
        left_margin = sect.left_margin
        right_margin = sect.right_margin
        printable_emu = page_width - left_margin - right_margin
        table.columns[0].width = Inches( (printable_emu / 914400.0) * 0.80 )
        table.columns[1].width = Inches( (printable_emu / 914400.0) * 0.20 )
        table.autofit = False
    except Exception:
        try:
            table.columns[0].width = Inches(5.5)
            table.columns[1].width = Inches(1.5)
            table.autofit = False
        except Exception:
            pass
    try:
        new_p = OxmlElement("w:p")
        table._element.addnext(new_p)
        after_para = Paragraph(new_p, getattr(ref, "_parent", doc))
        _zero_para_spacing(after_para)
        return after_para
    except Exception:
        spacer = _safe_add_paragraph_after(doc, table._element, text="")
        _zero_para_spacing(spacer)
        return spacer

def insert_certifications_block(doc, heading_paragraph, certs_list: List[Dict[str,Any]]):
    end_para = find_section_end_paragraph(doc, heading_paragraph) or heading_paragraph
    p_heading = _safe_add_paragraph_after(doc, end_para, text="")
    _zero_para_spacing(p_heading)
    try:
        run_h = p_heading.add_run("Certifications and Licenses")
        run_h.bold = True
        run_h.font.name = "Calibri"
        run_h.font.size = Pt(11)
    except Exception:
        p_heading.add_run("Certifications and Licenses")
    insert_ref = p_heading
    if not certs_list:
        spacer = _safe_add_paragraph_after(doc, insert_ref, text="")
        _zero_para_spacing(spacer)
        return spacer
    tbl_elm = OxmlElement("w:tbl")
    tbl_pr = OxmlElement("w:tblPr"); tbl_elm.append(tbl_pr)
    tbl_grid = OxmlElement("w:tblGrid")
    col1 = OxmlElement("w:gridCol"); col1.set(qn("w:w"), str(8000))
    col2 = OxmlElement("w:gridCol"); col2.set(qn("w:w"), str(1200))
    tbl_grid.append(col1); tbl_grid.append(col2)
    tbl_elm.append(tbl_grid)
    try:
        insert_ref._element.addnext(tbl_elm)
        parent = getattr(insert_ref, "_parent", None) or getattr(doc, "_element", None)
        table = _Table(tbl_elm, parent)
    except Exception:
        table = doc.add_table(rows=0, cols=2)
    try:
        tbl_root = table._tbl
        for tr in list(tbl_root.findall(qn('w:tr'))):
            tbl_root.remove(tr)
    except Exception:
        pass
    for cert in certs_list or []:
        if isinstance(cert, dict):
            name = (cert.get("name") or cert.get("certification") or cert.get("title") or "").strip()
            issuer = (cert.get("issuer") or cert.get("authority") or "").strip()
            date = (cert.get("date") or cert.get("issued") or cert.get("when") or "").strip()
            left_text = name
            if issuer:
                left_text += (" — " + issuer) if left_text else issuer
        else:
            left_text = str(cert).strip()
            date = ""
        row = table.add_row()
        left_cell = row.cells[0]; right_cell = row.cells[1]
        for p in list(left_cell.paragraphs):
            try: left_cell._element.remove(p._element)
            except Exception: pass
        for p in list(right_cell.paragraphs):
            try: right_cell._element.remove(p._element)
            except Exception: pass
        p_left = left_cell.add_paragraph()
        _zero_para_spacing(p_left)
        try:
            rleft = p_left.add_run("\u2022 " + left_text)
            rleft.font.name = "Calibri"
            rleft.font.size = Pt(11)
        except Exception:
            p_left.add_run(left_text)
        p_right = right_cell.add_paragraph()
        _zero_para_spacing(p_right)
        try:
            p_right.alignment = 2
        except Exception:
            pass
        try:
            rright = p_right.add_run(date or "")
            rright.font.name = "Calibri"
            rright.font.size = Pt(11)
        except Exception:
            p_right.add_run(date or "")
    try:
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for name in ('top','left','bottom','right','insideH','insideV'):
            b = OxmlElement(f'w:{name}')
            b.set(qn('w:val'), 'nil')
            tblBorders.append(b)
        tblPr.append(tblBorders)
    except Exception:
        pass
    try:
        sect = doc.sections[0]
        printable_emu = sect.page_width - sect.left_margin - sect.right_margin
        table.columns[0].width = Inches((printable_emu / 914400.0) * 0.80)
        table.columns[1].width = Inches((printable_emu / 914400.0) * 0.20)
        table.autofit = False
    except Exception:
        try:
            table.columns[0].width = Inches(5.5)
            table.columns[1].width = Inches(1.5)
            table.autofit = False
        except Exception:
            pass
    try:
        new_p = OxmlElement("w:p")
        table._element.addnext(new_p)
        after_para = Paragraph(new_p, getattr(insert_ref, "_parent", doc))
        _zero_para_spacing(after_para)
        return after_para
    except Exception:
        spacer = _safe_add_paragraph_after(doc, table._element, text="")
        _zero_para_spacing(spacer)
        return spacer

def insert_projects_block(doc, heading_paragraph, projects_list: List[Dict[str,Any]]):
    end_para = find_section_end_paragraph(doc, heading_paragraph) or heading_paragraph
    insert_ref = end_para
    p_heading = _safe_add_paragraph_after(doc, insert_ref, text="")
    _zero_para_spacing(p_heading)
    try:
        r = p_heading.add_run("Projects")
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)
    except Exception:
        p_heading.add_run("Projects")
    insert_ref = p_heading
    for proj in projects_list or []:
        title = (proj.get("title") or proj.get("name") or "").strip()
        org = (proj.get("organization") or proj.get("company") or "").strip()
        start = (proj.get("start_date") or "").strip()
        end = (proj.get("end_date") or "").strip()
        desc = proj.get("description") or proj.get("details") or []
        header_text = title
        if org:
            header_text += (" — " + org) if header_text else org
        if start or end:
            rng = (start + " – " + end).strip(" –")
            header_text += (" (" + rng + ")") if header_text else rng
        if header_text:
            p_hdr = _safe_add_paragraph_after(doc, insert_ref, text="")
            _zero_para_spacing(p_hdr)
            try:
                r = p_hdr.add_run(header_text)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
            except Exception:
                p_hdr.add_run(header_text)
            insert_ref = p_hdr
        desc_items = []
        if isinstance(desc, str):
            desc_items = [ln.strip() for ln in desc.splitlines() if ln.strip()] if "\n" in desc else ([desc.strip()] if desc.strip() else [])
        elif isinstance(desc, (list, tuple)):
            desc_items = [str(x).strip() for x in desc if x is not None and str(x).strip()]
        for d in desc_items:
            p_d = _safe_add_paragraph_after(doc, insert_ref, text="")
            _zero_para_spacing(p_d)
            try:
                p_d.style = "List Bullet"
                rr = p_d.add_run(d)
                rr.font.name = "Calibri"
                rr.font.size = Pt(11)
            except Exception:
                rr = p_d.add_run("\u2022 " + d)
                try:
                    rr.font.name = "Calibri"
                    rr.font.size = Pt(11)
                except Exception:
                    pass
            insert_ref = p_d
    spacer = _safe_add_paragraph_after(doc, insert_ref, text="")
    _zero_para_spacing(spacer)
    return spacer

def _process_paragraph_for_mapping(doc, paragraph, mapping: Dict[str, Any], bullet_style="List Bullet"):
    if paragraph is None:
        return False
    if not hasattr(doc, "_inserted_sections"):
        try:
            setattr(doc, "_inserted_sections", set())
        except Exception:
            doc._inserted_sections = set()
    inserted = getattr(doc, "_inserted_sections")
    original_text = paragraph.text or ""
    text = original_text
    changed = False
    def _is_heading_like(p_text: str, names: List[str]) -> bool:
        if not p_text:
            return False
        s = p_text.strip().lower()
        for nm in names:
            nm_l = nm.lower()
            if re.match(rf"^\s*{re.escape(nm_l)}\b", s):
                return True
        return False
    for k, v in mapping.items():
        if not k:
            continue
        if isinstance(v, (list, tuple, dict)):
            continue
        if k in text:
            try:
                text = text.replace(k, str(v))
                changed = True
            except Exception:
                logger.exception("Scalar replacement failed for key %s", k)
    for k, v in mapping.items():
        if not k or k not in text:
            continue
        if isinstance(v, dict):
            try:
                before, sep, after = text.partition(k)
                insert_ref = paragraph
                if before and before.strip():
                    insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=before)
                after_para = insert_skills_table(doc, insert_ref, v)
                if after and after.strip():
                    _safe_add_paragraph_after(doc, after_para, text=after)
                _remove_paragraph(paragraph)
                return True
            except Exception:
                logger.exception("Failed to expand dict placeholder %s", k)
    EMP_KEY = "{{EMPLOYMENT_HISTORY}}"
    if EMP_KEY in mapping and EMP_KEY in text:
        v = mapping[EMP_KEY]
        if isinstance(v, list) and _list_looks_like_experience_items(v):
            try:
                before, sep, after = text.partition(EMP_KEY)
                insert_ref = paragraph
                if before and before.strip():
                    insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=before)
                after_para = insert_experience_block(doc, insert_ref, v)
                if after and after.strip():
                    _safe_add_paragraph_after(doc, after_para, text=after)
                _remove_paragraph(paragraph)
                return True
            except Exception:
                logger.exception("Failed to expand employment placeholder %s", EMP_KEY)
    EDU_KEY = "{{EDUCATION}}"
    if EDU_KEY in mapping and EDU_KEY in text:
        v = mapping[EDU_KEY]
        if isinstance(v, (list, tuple)):
            try:
                before, sep, after = text.partition(EDU_KEY)
                insert_ref = paragraph
                if before and before.strip():
                    insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=before)
                after_para = insert_education_block(doc, insert_ref, v)
                if after and after.strip():
                    _safe_add_paragraph_after(doc, after_para, text=after)
                _remove_paragraph(paragraph)
                inserted.add("education")
                return True
            except Exception:
                logger.exception("Failed to expand education placeholder %s", EDU_KEY)
    CERT_KEY = "{{CERTIFICATIONS}}"
    if CERT_KEY in mapping and CERT_KEY in text:
        v = mapping[CERT_KEY]
        if isinstance(v, (list, tuple)):
            try:
                before, sep, after = text.partition(CERT_KEY)
                insert_ref = paragraph
                if before and before.strip():
                    insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=before)
                after_para = insert_certifications_block(doc, insert_ref, v)
                if after and after.strip():
                    _safe_add_paragraph_after(doc, after_para, text=after)
                _remove_paragraph(paragraph)
                inserted.add("certifications")
                return True
            except Exception:
                logger.exception("Failed to expand certifications placeholder %s", CERT_KEY)
    PROJ_KEY = "{{PROJECTS}}"
    if PROJ_KEY in mapping and PROJ_KEY in text:
        v = mapping[PROJ_KEY]
        if isinstance(v, (list, tuple)):
            try:
                before, sep, after = text.partition(PROJ_KEY)
                insert_ref = paragraph
                if before and before.strip():
                    insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=before)
                after_para = insert_projects_block(doc, insert_ref, v)
                if after and after.strip():
                    _safe_add_paragraph_after(doc, after_para, text=after)
                _remove_paragraph(paragraph)
                inserted.add("projects")
                return True
            except Exception:
                logger.exception("Failed to expand projects placeholder %s", PROJ_KEY)
    try:
        norm = (original_text or "").strip().lower()
        
        if (("experience" in norm or "employment" in norm or "work" in norm) or _is_heading_like(original_text, ["experience", "employment", "work"])) and "projects" not in inserted:
            proj_candidate = None
            if "{{PROJECTS}}" in mapping and isinstance(mapping["{{PROJECTS}}"], (list, tuple)):
                proj_candidate = mapping["{{PROJECTS}}"]
            elif "projects" in mapping and isinstance(mapping["projects"], (list, tuple)):
                proj_candidate = mapping["projects"]
            if proj_candidate:
                try:
                    insert_ref = paragraph
                    after_para = insert_projects_block(doc, insert_ref, proj_candidate)
                    inserted.add("projects")
                    return True
                except Exception:
                    logger.exception("Failed to auto-insert projects block after Employment heading")
    except Exception:
        logger.debug("Heading-detection for optional sections failed", exc_info=True)
    for k, v in mapping.items():
        if not k or k not in text:
            continue
        if isinstance(v, (list, tuple)):
            try:
                before, sep, after = text.partition(k)
                insert_ref = paragraph
                if before and before.strip():
                    insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=before)
                for item in v:
                    if isinstance(item, dict) and ("summary" in item or "bullets" in item):
                        bullets = item.get("summary") or item.get("bullets") or []
                        if isinstance(bullets, (list, tuple)):
                            for b in bullets:
                                p_b = _safe_add_paragraph_after(doc, insert_ref, text="")
                                _zero_para_spacing(p_b)
                                try:
                                    p_b.style = bullet_style
                                    run = p_b.add_run(str(b))
                                    try:
                                        run.font.name = "Calibri"
                                        run.font.size = Pt(11)
                                    except Exception:
                                        pass
                                except Exception:
                                    run = p_b.add_run("\u2022 " + str(b))
                                    try:
                                        run.font.name = "Calibri"
                                        run.font.size = Pt(11)
                                    except Exception:
                                        pass
                                insert_ref = p_b
                        else:
                            p_b = _safe_add_paragraph_after(doc, insert_ref, text="")
                            _zero_para_spacing(p_b)
                            try:
                                p_b.style = bullet_style
                                run = p_b.add_run(str(bullets))
                                try:
                                    run.font.name = "Calibri"
                                    run.font.size = Pt(11)
                                except Exception:
                                    pass
                            except Exception:
                                run = p_b.add_run("\u2022 " + str(bullets))
                                try:
                                    run.font.name = "Calibri"
                                    run.font.size = Pt(11)
                                except Exception:
                                    pass
                            insert_ref = p_b
                    else:
                        p_item = _safe_add_paragraph_after(doc, insert_ref, text="")
                        _zero_para_spacing(p_item)
                        try:
                            p_item.style = bullet_style
                            run = p_item.add_run(str(item))
                            try:
                                run.font.name = "Calibri"
                                run.font.size = Pt(11)
                            except Exception:
                                pass
                        except Exception:
                            run = p_item.add_run("\u2022 " + str(item))
                            try:
                                run.font.name = "Calibri"
                                run.font.size = Pt(11)
                            except Exception:
                                pass
                        insert_ref = p_item
                if after and after.strip():
                    insert_ref = _safe_add_paragraph_after(doc, insert_ref, text=after)
                _remove_paragraph(paragraph)
                return True
            except Exception:
                logger.exception("Failed to expand list placeholder %s", k)
    if changed and text != original_text:
        try:
            _replace_paragraph_text(paragraph, text)
            return True
        except Exception:
            logger.exception("Failed to replace paragraph text inline")
    return False

def replace_placeholders_in_doc(doc: Document, mapping: Dict[str, Any]):
    for p in list(doc.paragraphs):
        _process_paragraph_for_mapping(doc, p, mapping)
    for table in list(doc.tables):
        for row in list(table.rows):
            for cell in list(row.cells):
                for p in list(cell.paragraphs):
                    _process_paragraph_for_mapping(doc, p, mapping)
    for section in doc.sections:
        header = section.header
        for p in list(header.paragraphs):
            _process_paragraph_for_mapping(doc, p, mapping)
        footer = section.footer
        for p in list(footer.paragraphs):
            _process_paragraph_for_mapping(doc, p, mapping)

def build_employment_text(employment_list: List[Dict[str,Any]]) -> str:
    parts = []
    for j in employment_list or []:
        header = []
        if j.get("position"):
            header.append(j.get("position"))
        if j.get("company"):
            header.append("at " + j.get("company"))
        header_line = " ".join(header)
        date_line = ""
        if j.get("location") or j.get("start_date"):
            date_line = f"{j.get('location','')} | {j.get('start_date','')} – {j.get('end_date','')}".strip()
        bullets = j.get("summary",[]) or []
        btext = "\n".join([f"• {b}" for b in bullets])
        chunk = header_line
        if date_line:
            chunk += "\n" + date_line
        if btext:
            chunk += "\n" + btext
        parts.append(chunk)
    return "\n\n".join(parts)

def build_education_text(education_list: List[Dict[str,Any]]) -> str:
    return "\n".join([f"{e.get('degree','')} — {e.get('institution','')} ({e.get('end_year','')})" for e in education_list or []])

def insert_dynamic_table(doc, section_placeholder, data_list, headers=None):
    # Find the section header paragraph where we want to insert the table
    section_para = _find_heading_paragraph(doc, [section_placeholder])
    if not section_para:
        return  # No heading found, don't insert anything

    # Create a table with the correct number of columns based on the headers
    table_elm = OxmlElement("w:tbl")
    tbl_pr = OxmlElement("w:tblPr")
    table_elm.append(tbl_pr)

    tbl_grid = OxmlElement("w:tblGrid")
    if headers:
        for _ in range(len(headers)):
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), "2400")  # Adjust width as needed
            tbl_grid.append(grid_col)
    table_elm.append(tbl_grid)

    # Add table after the section header
    section_para._element.addnext(table_elm)

    # Create a Table object
    parent = getattr(section_para, "_parent", None) or getattr(doc, "_element", None)
    table = _Table(table_elm, parent)

    # Add headers if needed
    if headers:
        header_row = table.add_row()
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.paragraphs[0].add_run(header).bold = True

    # Insert data into the table rows
    for data in data_list:
        row = table.add_row()
        for i, key in enumerate(headers):
            value = data.get(key.lower(), "")  # Assuming dictionary with keys matching headers
            cell = row.cells[i]
            cell.paragraphs[0].add_run(str(value))

    return table

# ============================================================
# OCFO ADVANCED REPLACEMENT HELPERS (must be defined BEFORE use)
# ============================================================

from docx.shared import RGBColor

def _replace_scalar_across_runs(paragraph, placeholder, value, name_style=False):
    if paragraph is None:
        return False
    runs = list(getattr(paragraph, "runs", []))
    if not runs:
        return False

    joined = "".join([r.text or "" for r in runs])
    if placeholder not in joined:
        return False

    new_text = joined.replace(placeholder, str(value))

    try:
        for r in list(paragraph.runs):
            paragraph._element.remove(r._element)
    except Exception:
        pass

    new_run = paragraph.add_run(new_text)

    if name_style and value:
        try:
            new_run.font.name = "Times New Roman"
            new_run.font.size = Pt(18)
            new_run.bold = True
            new_run.font.color.rgb = RGBColor(0, 0, 255)
        except Exception:
            pass
    else:
        try:
            new_run.font.name = "Calibri"
            new_run.font.size = Pt(11)
        except Exception:
            pass

    return True


def _aggressive_scalar_replace(doc, mapping):
    scalar_keys = [k for k, v in mapping.items() if not isinstance(v, (list, tuple, dict))]

    for p in list(doc.paragraphs):
        for k in scalar_keys:
            if k in (p.text or "") or any(k in (r.text or "") for r in p.runs):
                is_name = k in ("{{CANDIDATE_NAME}}", "{{FULL_NAME}}")
                _replace_scalar_across_runs(p, k, mapping[k], name_style=is_name)

    for table in list(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    for k in scalar_keys:
                        if k in (p.text or "") or any(k in (r.text or "") for r in p.runs):
                            is_name = k in ("{{CANDIDATE_NAME}}", "{{FULL_NAME}}")
                            _replace_scalar_across_runs(p, k, mapping[k], name_style=is_name)

    for section in doc.sections:
        for p in list(section.header.paragraphs):
            for k in scalar_keys:
                if k in (p.text or ""):
                    is_name = k in ("{{CANDIDATE_NAME}}", "{{FULL_NAME}}")
                    _replace_scalar_across_runs(p, k, mapping[k], name_style=is_name)

        for p in list(section.footer.paragraphs):
            for k in scalar_keys:
                if k in (p.text or ""):
                    is_name = k in ("{{CANDIDATE_NAME}}", "{{FULL_NAME}}")
                    _replace_scalar_across_runs(p, k, mapping[k], name_style=is_name)


def _find_heading_paragraph(doc, headings):
    for p in list(doc.paragraphs):
        txt = (p.text or "").strip().lower()
        for h in headings:
            if re.match(rf"^\s*{re.escape(h)}\b", txt, flags=re.IGNORECASE):
                return p
    return None

def create_styled_table(doc, ref_paragraph, headers, rows):
    insert_ref = _safe_add_paragraph_after(doc, ref_paragraph, text="")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = None
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        for p in list(cell.paragraphs):
            try:
                cell._element.remove(p._element)
            except Exception:
                pass
        p = cell.add_paragraph()
        run = p.add_run(str(h))
        try:
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass
        _zero_para_spacing(p)
    for rdata in rows or []:
        row = table.add_row()
        for i, key in enumerate(headers):
            value = ""
            if isinstance(rdata, dict):
                lookup_key = key.lower().strip()
                lookup_key = lookup_key.replace("(", "").replace(")", "").replace("/", "_").replace(" ", "_").replace(":", "")
                if lookup_key in rdata:
                    value = rdata.get(lookup_key) or ""
                else:
                    alt_val = ""
                    for alt in ("degree","area_of_study","area","major","institution","school","location","awarded","end_year","date","start_year","certification","name","issuer","issued_by","date_obtained","certification_number","number","expiration_date","expires"):
                        if alt in rdata and (alt in lookup_key or lookup_key in alt or alt.replace("_","") in lookup_key.replace("_","")):
                            alt_val = rdata.get(alt) or ""
                            break
                    if not alt_val:
                        for k,v in rdata.items():
                            if isinstance(v, str) and len(v) > 0:
                                alt_val = v
                                break
                    value = alt_val or ""
            else:
                value = str(rdata)
            cell = row.cells[i]
            for p in list(cell.paragraphs):
                try:
                    cell._element.remove(p._element)
                except Exception:
                    pass
            p = cell.add_paragraph()
            run = p.add_run(str(value or ""))
            try:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass
            _zero_para_spacing(p)
    try:
        sect = doc.sections[0]
        printable_emu = sect.page_width - sect.left_margin - sect.right_margin
        total_inches = printable_emu / 914400.0
        ncols = len(headers)
        col_width = total_inches / ncols if ncols else total_inches
        for i in range(ncols):
            try:
                table.columns[i].width = Inches(col_width)
            except Exception:
                pass
        table.autofit = False
    except Exception:
        try:
            table.autofit = True
        except Exception:
            pass
    return table

def build_education_table(doc, heading_paragraph, education_list):
    headers = [
        "Degree (AA/AS, BA/BS, MA/MS/MBA, PhD, JD)",
        "Area of Study",
        "School/College/University",
        "Location",
        "Was the degree awarded? (Yes/No)",
        "OPTIONAL: Date (MM/YY)"
    ]
    rows = []
    for edu in education_list or []:
        row = {}
        row_key_map = {}
        row_key_map[headers[0].lower().replace(" ", "_")] = edu.get("degree","") or edu.get("degree","")
        row_key_map[headers[1].lower().replace(" ", "_")] = edu.get("area_of_study","") or edu.get("major","") or edu.get("area","")
        row_key_map[headers[2].lower().replace(" ", "_")] = edu.get("institution","") or edu.get("school","")
        row_key_map[headers[3].lower().replace(" ", "_")] = edu.get("location","")
        row_key_map[headers[4].lower().replace(" ", "_")] = edu.get("awarded","") or edu.get("was_awarded","") or ""
        date_val = edu.get("end_year","") or edu.get("date","") or edu.get("start_year","") or edu.get("year","")
        row_key_map[headers[5].lower().replace(" ", "_")] = date_val
        rows.append(row_key_map)
    return create_styled_table(doc, heading_paragraph, headers, rows)

def build_certifications_table(doc, heading_paragraph, certs_list):
    headers = [
        "Certification",
        "Issued By",
        "Date Obtained (MM/YY)",
        "Certification Number (If Applicable)",
        "Expiration Date (If Applicable)"
    ]
    rows = []
    for c in certs_list or []:
        row = {}
        name = ""
        issuer = ""
        date = ""
        number = ""
        exp = ""
        if isinstance(c, dict):
            name = c.get("name") or c.get("certification") or c.get("title") or c.get("certificate") or c.get("certification_name") or ""
            issuer = c.get("issuer") or c.get("authority") or c.get("issued_by") or c.get("issuer_name") or ""
            date = c.get("date") or c.get("issued") or c.get("date_obtained") or ""
            number = c.get("number") or c.get("certification_number") or c.get("id") or c.get("cert_id") or ""
            exp = c.get("expiration_date") or c.get("expires") or c.get("expiry") or ""
        else:
            name = str(c)
        row[headers[0].lower().replace(" ", "_")] = name
        row[headers[1].lower().replace(" ", "_")] = issuer
        row[headers[2].lower().replace(" ", "_")] = date
        row[headers[3].lower().replace(" ", "_")] = number
        row[headers[4].lower().replace(" ", "_")] = exp
        rows.append(row)
    return create_styled_table(doc, heading_paragraph, headers, rows)

def adjust_paragraph_spacing(doc):
    for paragraph in doc.paragraphs:
        # Remove space after the paragraph
        paragraph.paragraph_format.space_after = 0
        # Remove space before the paragraph
        paragraph.paragraph_format.space_before = 0
        # Ensure no extra spacing within the paragraph
        paragraph.paragraph_format.line_spacing = 1.0  # Single line spacing
    return doc

@app.post("/format_from_template")  
async def format_from_template(
    resume_json: UploadFile = File(..., description="Extracted resume JSON"),
    template_name: str = Form(..., description="Template file name in templates directory")
):
    try:
        raw = await resume_json.read()
        data = json.loads(raw.decode("utf-8"))

        education_data = data.get("education", []) or []
        certifications_data = data.get("certifications", []) or []

        template_path = os.path.join(TEMPLATES_DIR, template_name)
        if not os.path.exists(template_path):
            return JSONResponse(status_code=404, content={"error": f"Template '{template_name}' not found."})

        if template_path.lower().endswith(".doc"):
            template_path = convert_doc_to_docx(template_path)

        doc = Document(template_path)
        
        doc = adjust_paragraph_spacing(doc)

        candidate = data.get("candidate", {}) or {}

        inferred_title = ""
        if candidate.get("title") or candidate.get("position") or candidate.get("role"):
            inferred_title = candidate.get("title") or candidate.get("position") or candidate.get("role")
        else:
            eh = data.get("employment_history", []) or []
            if isinstance(eh, list) and len(eh) > 0 and isinstance(eh[0], dict):
                inferred_title = eh[0].get("position") or eh[0].get("title") or eh[0].get("role") or ""

        vector_vms = data.get("vector_vms_requisition_number") or data.get("requisition_number") or data.get("vms_requisition") or data.get("requisition") or ""

        placeholder_map = {
            "{{CANDIDATE_NAME}}": candidate.get("full_name", "") or "",
            "{{FULL_NAME}}": candidate.get("full_name", "") or "",
            "{{EMAIL}}": candidate.get("email", "") or "",
            "{{PHONE}}": candidate.get("phone", "") or "",
            "{{ADDRESS}}": candidate.get("address", "") or "",
            "{{LINKEDIN}}": candidate.get("linkedin", "") or "",
            "{{SUMMARY}}": data.get("professional_summary", "") or "",
            "{{SKILLS}}": data.get("skills", {}) or {},
            "{{EMPLOYMENT_HISTORY}}": data.get("employment_history", []) or [],
            "{{EDUCATION}}": education_data or [],
            "{{CERTIFICATIONS}}": certifications_data or [],
            "{{PROJECTS}}": data.get("projects", []) or [],
            "{{TITLE_ROLE}}": inferred_title or "",
            "{{REQUISITION}}": vector_vms or "",
            "{{VECTOR_VMS_REQUISITION_NUMBER}}": vector_vms or ""
        }

        audit_base = os.path.splitext(template_name)[0]
        debug_map_path = os.path.join(OUTPUT_DIR, f"{audit_base}_placeholder_map.json")
        with open(debug_map_path, "w", encoding="utf-8") as f:
            json.dump(placeholder_map, f, indent=2, ensure_ascii=False)

        doc_before_path = os.path.join(OUTPUT_DIR, f"{audit_base}_before_fill.docx")
        doc.save(doc_before_path)
        
        def _find_heading_paragraph(doc_obj, headings):
            for p in list(doc_obj.paragraphs):
                txt = (p.text or "").strip().lower()
                for h in headings:
                    if re.match(rf"^\s*{re.escape(h)}\b", txt, flags=re.IGNORECASE):
                        return p
            return None

        def _set_paragraph_font(paragraph, font_name, size_pt, bold=False, color_rgb=(0,0,0)):
            for r in list(getattr(paragraph, "runs", [])):
                try:
                    r.font.name = font_name
                    r.font.size = Pt(size_pt)
                    r.font.bold = bold
                    r.font.color.rgb = RGBColor(*color_rgb)
                except Exception:
                    pass

        def _make_table_borders(tbl):
            try:
                tbl_pr = tbl._element.find(qn('w:tblPr'))
                if tbl_pr is None:
                    tbl_pr = OxmlElement('w:tblPr')
                    tbl._element.insert(0, tbl_pr)
            except Exception:
                tbl_pr = tbl._element.find(qn('w:tblPr'))
                if tbl_pr is None:
                    tbl_pr = OxmlElement('w:tblPr')
                    tbl._element.insert(0, tbl_pr)
            tbl_borders = tbl_pr.find(qn('w:tblBorders'))
            if tbl_borders is not None:
                try:
                    tbl_pr.remove(tbl_borders)
                except Exception:
                    pass
            tbl_borders = OxmlElement('w:tblBorders')
            for name in ('top','left','bottom','right','insideH','insideV'):
                b = OxmlElement(f'w:{name}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), '000000')
                tbl_borders.append(b)
            tbl_pr.append(tbl_borders)

        def insert_education_table(document, heading_paragraph, edu_list):
            if heading_paragraph is None:
                return
            h = heading_paragraph
            insert_after = h
            before, sep, after = (h.text or "").partition("{{EDUCATION}}") if "{{EDUCATION}}" in (h.text or "") else (None, None, None)
            if before and before.strip():
                insert_after = _safe_add_paragraph_after(document, insert_after, text=before)
            cols = [
                "Degree (AA/AS, BA/BS, MA/MS/MBA, PhD, JD)",
                "Area of Study",
                "School/College/University",
                "Location",
                "Was the degree awarded? (Yes/No)",
                "OPTIONAL: Date (MM/YY)"
            ]
            tbl_elm = OxmlElement("w:tbl")
            tbl_pr = OxmlElement("w:tblPr"); tbl_elm.append(tbl_pr)
            try:
                sect = document.sections[0]
                printable_emu = sect.page_width - sect.left_margin - sect.right_margin
                total_inches = printable_emu / 914400.0
                printable_twips = int(total_inches * 1440)
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), str(printable_twips))
                tblW.set(qn('w:type'), 'dxa')
                tbl_pr.append(tblW)
            except Exception:
                total_inches = 6.75
                printable_twips = int(total_inches * 1440)
            col_widths = [0.18, 0.18, 0.34, 0.12, 0.09, 0.09]
            tbl_grid = OxmlElement("w:tblGrid")
            for cw in col_widths:
                col_twips = int(printable_twips * float(cw))
                g = OxmlElement("w:gridCol"); g.set(qn("w:w"), str(col_twips)); tbl_grid.append(g)
            tbl_elm.append(tbl_grid)
            insert_after._element.addnext(tbl_elm)
            parent = getattr(insert_after, "_parent", None) or getattr(document, "_element", None)
            table = _Table(tbl_elm, parent)
            hdr_row = table.add_row()
            for i, label in enumerate(cols):
                cell = hdr_row.cells[i]
                for p in list(cell.paragraphs):
                    try:
                        cell._element.remove(p._element)
                    except Exception:
                        pass
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(label)
                try:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
                    r.bold = True
                    r.font.color.rgb = RGBColor(0,0,0)
                except Exception:
                    pass
            for edu in edu_list or []:
                degree = (edu.get("degree") or "").strip()
                area = (edu.get("area_of_study") or edu.get("major") or "").strip()
                inst = (edu.get("institution") or edu.get("school") or "").strip()
                loc = (edu.get("location") or "").strip()
                awarded = "Yes" if str(edu.get("awarded") or "").strip().lower() in ("yes","y","true","1") else (str(edu.get("awarded") or "").strip() or "")
                date = (edu.get("end_year") or edu.get("date") or edu.get("start_year") or "").strip()
                row = table.add_row()
                vals = [degree, area, inst, loc, awarded, date]
                for i, val in enumerate(vals):
                    cell = row.cells[i]
                    for p in list(cell.paragraphs):
                        try:
                            cell._element.remove(p._element)
                        except Exception:
                            pass
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = p.add_run(val or "")
                    try:
                        if val.strip() in ["(AA/AS, BA/BS, MA/MS/MBA, PhD, JD)", "(Yes/No)", "(MM/YY)", "(If Applicable)"]:
                            run.font.bold = False
                        else:
                            run.font.bold = True
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0,0,0)
                        run.font.bold = False
                    except Exception:
                        pass
            try:
                sect = document.sections[0]
                printable_emu = sect.page_width - sect.left_margin - sect.right_margin
                total_inches = printable_emu / 914400.0
                col_widths_local = [0.18, 0.18, 0.34, 0.12, 0.09, 0.09]
                for i, cw in enumerate(col_widths_local):
                    try:
                        table.columns[i].width = Inches(total_inches * cw)
                    except Exception:
                        pass
                table.autofit = False
            except Exception:
                pass
            _make_table_borders(table)
            if after and after.strip():
                _safe_add_paragraph_after(document, table._element, text=after)
            return table

        def insert_certifications_table(document, heading_paragraph, certs_list):
            if heading_paragraph is None:
                return
            h = heading_paragraph
            insert_after = h
            before, sep, after = (h.text or "").partition("{{CERTIFICATIONS}}") if "{{CERTIFICATIONS}}" in (h.text or "") else (None, None, None)
            if before and before.strip():
                insert_after = _safe_add_paragraph_after(document, insert_after, text=before)
            cols = [
                "Certification",
                "Issued By",
                "Date Obtained (MM/YY)",
                "Certification Number (If Applicable)",
                "Expiration Date (If Applicable)"
            ]
            tbl_elm = OxmlElement("w:tbl")
            tbl_pr = OxmlElement("w:tblPr"); tbl_elm.append(tbl_pr)
            try:
                sect = document.sections[0]
                printable_emu = sect.page_width - sect.left_margin - sect.right_margin
                total_inches = printable_emu / 914400.0
                printable_twips = int(total_inches * 1440)
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), str(printable_twips))
                tblW.set(qn('w:type'), 'dxa')
                tbl_pr.append(tblW)
            except Exception:
                total_inches = 6.75
                printable_twips = int(total_inches * 1440)
            col_widths = [0.30, 0.25, 0.12, 0.16, 0.17]
            tbl_grid = OxmlElement("w:tblGrid")
            for cw in col_widths:
                col_twips = int(printable_twips * float(cw))
                g = OxmlElement("w:gridCol"); g.set(qn("w:w"), str(col_twips)); tbl_grid.append(g)
            tbl_elm.append(tbl_grid)
            insert_after._element.addnext(tbl_elm)
            parent = getattr(insert_after, "_parent", None) or getattr(document, "_element", None)
            table = _Table(tbl_elm, parent)
            hdr_row = table.add_row()
            for i, label in enumerate(cols):
                cell = hdr_row.cells[i]
                for p in list(cell.paragraphs):
                    try:
                        cell._element.remove(p._element)
                    except Exception:
                        pass
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(label)
                try:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
                    r.bold = True
                    r.font.color.rgb = RGBColor(0,0,0)
                except Exception:
                    pass
            for cert in certs_list or []:
                name = (cert.get("name") or cert.get("certification") or cert.get("title") or "").strip()
                issuer = (cert.get("issuer") or cert.get("issued_by") or cert.get("authority") or "").strip()
                date = (cert.get("date") or cert.get("date_obtained") or cert.get("issued") or "").strip()
                number = (cert.get("certification_number") or cert.get("number") or "").strip()
                exp = (cert.get("expiration_date") or cert.get("expiry") or cert.get("expires") or "").strip()
                row = table.add_row()
                vals = [name, issuer, date, number, exp]
                for i, val in enumerate(vals):
                    cell = row.cells[i]
                    for p in list(cell.paragraphs):
                        try:
                            cell._element.remove(p._element)
                        except Exception:
                            pass
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    run = p.add_run(val or "")
                    try:
                        if val.strip() in ["(Yes/No)", "(MM/YY)", "(If Applicable)"]:
                            run.font.bold = False
                        else:
                            run.font.bold = True
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0,0,0)
                        run.font.bold = False
                    except Exception:
                        pass
            try:
                sect = document.sections[0]
                printable_emu = sect.page_width - sect.left_margin - sect.right_margin
                total_inches = printable_emu / 914400.0
                col_widths_local = [0.30, 0.25, 0.12, 0.16, 0.17]
                for i, cw in enumerate(col_widths_local):
                    try:
                        table.columns[i].width = Inches(total_inches * cw)
                    except Exception:
                        pass
                table.autofit = False
            except Exception:
                pass
            _make_table_borders(table)
            if after and after.strip():
                _safe_add_paragraph_after(document, table._element, text=after)
            return table

        placeholder_paragraphs = {}
        all_placeholders = list(placeholder_map.keys())
        for p in list(doc.paragraphs):
            txt = p.text or ""
            for k in all_placeholders:
                if k in txt:
                    placeholder_paragraphs.setdefault(k, []).append(p)
        for table in list(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for p in list(cell.paragraphs):
                        txt = p.text or ""
                        for k in all_placeholders:
                            if k in txt:
                                placeholder_paragraphs.setdefault(k, []).append(p)
        for section in doc.sections:
            header = section.header
            for p in list(header.paragraphs):
                txt = p.text or ""
                for k in all_placeholders:
                    if k in txt:
                        placeholder_paragraphs.setdefault(k, []).append(p)
            footer = section.footer
            for p in list(footer.paragraphs):
                txt = p.text or ""
                for k in all_placeholders:
                    if k in txt:
                        placeholder_paragraphs.setdefault(k, []).append(p)

        orig_para_texts = {}
        for k, paras in placeholder_paragraphs.items():
            for p in paras:
                try:
                    orig_para_texts[p] = p.text or ""
                except Exception:
                    pass

        try:
            _aggressive_scalar_replace(doc, placeholder_map)
        except Exception:
            logger.exception("Aggressive scalar replacement failed")

        try:
            replace_placeholders_in_doc(doc, placeholder_map)
        except Exception:
            logger.exception("Primary placeholder replacement failed")

        all_text = "\n".join([p.text or "" for p in doc.paragraphs])
        remaining = re.findall(r"\{\{\s*([^}]+?)\s*\}\}", all_text)

        if remaining:
            fallback_map = {}
            for ph in remaining:
                k_braced = "{{" + ph + "}}"
                if k_braced in placeholder_map:
                    fallback_map[k_braced] = placeholder_map[k_braced]
                else:
                    key_plain = ph.strip()
                    if key_plain in placeholder_map:
                        fallback_map["{{" + key_plain + "}}"] = placeholder_map[key_plain]
                    elif key_plain.upper() in [kk.strip("{}").upper() for kk in placeholder_map]:
                        matches = [kk for kk in placeholder_map if kk.strip("{}").upper() == key_plain.upper()]
                        if matches:
                            fallback_map["{{" + key_plain + "}}"] = placeholder_map[matches[0]]
            if fallback_map:
                try:
                    _aggressive_scalar_replace(doc, fallback_map)
                except Exception:
                    logger.exception("Fallback aggressive replace failed")
                try:
                    replace_placeholders_in_doc(doc, fallback_map)
                except Exception:
                    logger.exception("Fallback placeholder replacement failed")

        dark_blue = RGBColor(0, 51, 102)

        label_patterns = [
            re.compile(r'^\s*(Title\s*/\s*Role\s*:)', flags=re.IGNORECASE),
            re.compile(r'^\s*(Vector\s+VMS\s+Requisition\s+Number\s*:)', flags=re.IGNORECASE),
            re.compile(r'^\s*(Vector\s+VMS\s+Requisition\s*:)', flags=re.IGNORECASE)
        ]

        def apply_label_value_styling(paragraph):
            txt = paragraph.text or ""
            stripped = txt.strip()
            if not stripped:
                return
            label_match = None
            for pat in label_patterns:
                m = pat.match(stripped)
                if m:
                    label_match = m
                    break
            if not label_match:
                return
            label_text = label_match.group(1)
            remainder = stripped[label_match.end():].lstrip()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if not remainder:
                raw = paragraph.text or ""
                for ph in ("{{TITLE_ROLE}}", "{{REQUISITION}}", "{{VECTOR_VMS_REQUISITION_NUMBER}}"):
                    if ph in raw:
                        remainder = placeholder_map.get(ph, "") or ""
                        break
            try:
                for r in list(getattr(paragraph, "runs", [])):
                    try:
                        paragraph._element.remove(r._element)
                    except Exception:
                        pass
                run_label = paragraph.add_run(label_text)
                try:
                    run_label.font.name = "Times New Roman"
                    run_label.font.size = Pt(14)
                    run_label.bold = True
                    run_label.font.color.rgb = dark_blue
                except Exception:
                    pass
                paragraph.add_run(" ")
                value_text = remainder
                if not value_text:
                    if "{{TITLE_ROLE}}" in placeholder_map and placeholder_map.get("{{TITLE_ROLE}}"):
                        value_text = placeholder_map.get("{{TITLE_ROLE}}")
                    elif "{{REQUISITION}}" in placeholder_map and placeholder_map.get("{{REQUISITION}}"):
                        value_text = placeholder_map.get("{{REQUISITION}}")
                    elif "{{VECTOR_VMS_REQUISITION_NUMBER}}" in placeholder_map and placeholder_map.get("{{VECTOR_VMS_REQUISITION_NUMBER}}"):
                        value_text = placeholder_map.get("{{VECTOR_VMS_REQUISITION_NUMBER}}")
                run_val = paragraph.add_run(value_text or "")
                try:
                    run_val.font.name = "Calibri"
                    run_val.font.size = Pt(11)
                    run_val.bold = False
                    run_val.font.color.rgb = RGBColor(0,0,0)
                except Exception:
                    pass
            except Exception:
                logger.exception("Failed to apply label/value styling")

        for p in list(doc.paragraphs):
            apply_label_value_styling(p)

        for table in list(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for p in list(cell.paragraphs):
                        apply_label_value_styling(p)

        for section in doc.sections:
            for p in list(section.header.paragraphs):
                apply_label_value_styling(p)
            for p in list(section.footer.paragraphs):
                apply_label_value_styling(p)

        edu_heading = _find_heading_paragraph(doc, ["education", "education:"])
        cert_heading = _find_heading_paragraph(doc, ["certifications", "certifications:", "certifications and licenses", "certification", "certification:"])

        if edu_heading:
            insert_education_table(doc, edu_heading, education_data)

        if cert_heading:
            insert_certifications_table(doc, cert_heading, certifications_data)

        doc_after_path = os.path.join(OUTPUT_DIR, f"{audit_base}_after_fill.docx")
        doc.save(doc_after_path)

        safe_name = (candidate.get("full_name") or "candidate").strip().replace(" ", "_")
        out_filename = f"formatted_{audit_base}_{safe_name}.docx"
        out_filename = re.sub(r"[^\w\-_\.]", "_", out_filename)

        out_path = os.path.join(OUTPUT_DIR, out_filename)
        doc.save(out_path)

        return FileResponse(out_path, filename=out_filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        logger.exception("Fatal error in format_from_template")
        return JSONResponse(status_code=500, content={"error": str(e)})
    
@app.get("/download/{filename}")
def download_file(filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        return JSONResponse(status_code=404, content={"error": "File not found"})
    if filename.lower().endswith(".html"):
        with open(path, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read(), status_code=200)
    return FileResponse(path, filename=filename)

@app.get("/health")
def health():
    return {
        "message": "Smart Resume Formatter (Ollama Cloud)",
        "templates_dir": TEMPLATES_DIR,
        "output_dir": OUTPUT_DIR,
        "ollama_host": OLLAMA_API_HOST,
        "model": CLOUD_MODEL,
        "ollama_api_key_present": bool(OLLAMA_API_KEY)
    }

@app.get("/")
def root():
    return {"message": "Smart Resume Formatter running. Use /docs for interactive UI."}
